"""
后端服务：接收 docx 文件和批注数据，使用 python-docx 添加批注
"""
from flask import Flask, request, send_file
from flask_cors import CORS
from docx import Document
import io
import json
from datetime import datetime

app = Flask(__name__)
CORS(app)  # 允许跨域请求


def find_text_in_runs(paragraph, text_to_find):
    """
    在段落的 runs 中查找文本，返回匹配的 runs 列表
    """
    full_text = paragraph.text
    if text_to_find not in full_text:
        return None
    
    start_idx = full_text.find(text_to_find)
    end_idx = start_idx + len(text_to_find)
    
    # 计算文本在 runs 中的位置
    current_pos = 0
    matching_runs = []
    
    for run in paragraph.runs:
        run_len = len(run.text)
        run_start = current_pos
        run_end = current_pos + run_len
        
        # 检查这个 run 是否与目标文本重叠
        if run_start < end_idx and run_end > start_idx:
            matching_runs.append(run)
        
        current_pos += run_len
        
        # 如果已经超过结束位置，可以提前退出
        if current_pos >= end_idx:
            break
    
    return matching_runs if matching_runs else None


def add_comment_to_document(doc, comment_data):
    """
    在文档中添加批注
    comment_data 格式：
    {
        "type": "location" | "missing_clause" | "manual",
        "location": "要查找的文本",
        "fixed": true/false,
        "riskIndex": 1,
        "riskDesc": "风险说明",
        "suggestion": "修改意见",
        "content": "批注内容（手动批注）",
        "userName": "用户名"
    }
    """
    # 构建批注内容
    if comment_data['type'] == 'manual':
        comment_text = comment_data.get('content', '')
        author = comment_data.get('userName', '用户')
    else:
        fixed_status = '【已修复】' if comment_data.get('fixed', False) else '【未修复】'
        
        # 支持多个风险序号和描述
        risk_indices = comment_data.get('riskIndices', [])
        risk_descs = comment_data.get('riskDescs', [])
        
        # 向后兼容：如果没有 riskIndices，使用 riskIndex
        if not risk_indices:
            risk_index = comment_data.get('riskIndex', 0)
            if risk_index > 0:
                risk_indices = [risk_index]
        
        # 向后兼容：如果没有 riskDescs，使用 riskDesc
        if not risk_descs:
            risk_desc = comment_data.get('riskDesc', '')
            if risk_desc:
                risk_descs = [risk_desc]
        
        # 构建序号字符串
        if risk_indices:
            indices_str = '、'.join([f'#{idx}' for idx in risk_indices])
            indices_line = f"序号：{indices_str}"
        else:
            indices_line = ""
        
        # 构建风险说明字符串
        if risk_descs:
            risk_desc_str = '；'.join(risk_descs)
            risk_desc_line = f"风险说明：{risk_desc_str}"
        else:
            risk_desc_line = ""
        
        suggestion = comment_data.get('suggestion', '')
        
        # 组合批注内容
        parts = [fixed_status]
        if indices_line:
            parts.append(indices_line)
        if risk_desc_line:
            parts.append(risk_desc_line)
        if suggestion:
            parts.append(f"修改意见：{suggestion}")
        
        comment_text = '\n'.join(parts)
        author = 'SCAi Review'
    
    # 查找要添加批注的文本
    location = comment_data.get('location', '')
    if not location:
        return False
    
    if location == 'title':
        # 如果是标题，查找第一个段落
        if doc.paragraphs:
            paragraph = doc.paragraphs[0]
            runs = paragraph.runs
            if runs:
                try:
                    doc.add_comment(runs=runs, text=comment_text, author=author)
                    return True
                except Exception as e:
                    print(f"添加批注到标题时出错: {e}")
                    import traceback
                    traceback.print_exc()
        return False
    
    # 在文档中查找文本（包括所有段落和表格）
    # 先查找普通段落
    for paragraph in doc.paragraphs:
        if location in paragraph.text:
            # 查找匹配的 runs
            matching_runs = find_text_in_runs(paragraph, location)
            if matching_runs:
                try:
                    # 使用 python-docx 的 add_comment 方法
                    doc.add_comment(runs=matching_runs, text=comment_text, author=author)
                    return True
                except Exception as e:
                    print(f"添加批注时出错: {e}")
                    import traceback
                    traceback.print_exc()
                    # 继续查找其他位置
    
    # 如果没找到，也在表格中查找
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if location in paragraph.text:
                        matching_runs = find_text_in_runs(paragraph, location)
                        if matching_runs:
                            try:
                                doc.add_comment(runs=matching_runs, text=comment_text, author=author)
                                return True
                            except Exception as e:
                                print(f"在表格中添加批注时出错: {e}")
                                import traceback
                                traceback.print_exc()
    
    print(f"警告：未找到文本 '{location}'，无法添加批注")
    return False


@app.route('/api/add-comments', methods=['POST'])
def add_comments():
    """
    接收 docx 文件和批注数据，添加批注后返回
    """
    try:
        # 检查文件是否存在
        if 'docx' not in request.files:
            return {'error': '未找到 docx 文件'}, 400
        
        # 获取文件
        docx_file = request.files['docx']
        
        # 获取批注数据
        comments_json = request.form.get('comments', '[]')
        comments_data = json.loads(comments_json)
        
        # 读取 docx 文件
        doc = Document(io.BytesIO(docx_file.read()))
        
        # 添加批注
        for comment_data in comments_data:
            add_comment_to_document(doc, comment_data)
        
        # 保存到内存
        output = io.BytesIO()
        doc.save(output)
        output.seek(0)
        
        # 返回修改后的文档
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            as_attachment=True,
            download_name=docx_file.filename or 'document.docx'
        )
        
    except Exception as e:
        print(f"处理请求时出错: {e}")
        import traceback
        traceback.print_exc()
        return {'error': str(e)}, 500




if __name__ == '__main__':
    app.run(debug=True, port=8000, host='0.0.0.0')
